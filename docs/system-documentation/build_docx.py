# -*- coding: utf-8 -*-
"""
Build DRAMS_Live_System_Documentation.docx from rendered Mermaid PNGs + narrative.
Run after mermaid-cli has produced diagrams/render-*.png
"""
import os, glob, re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
DIAG = os.path.join(HERE, "diagrams")

# ---- locate rendered diagrams in order (render-1.png, render-2.png, ...) ----
def diagram_list():
    files = glob.glob(os.path.join(DIAG, "render-*.png"))
    def num(f):
        m = re.search(r'render-(\d+)\.png$', f)
        return int(m.group(1)) if m else 9999
    return sorted(files, key=num)

DIAGRAMS = diagram_list()
print("Found %d rendered diagrams" % len(DIAGRAMS))

# Map: ordinal index (0-based, as they appear in the markdown) -> caption
CAPTIONS = [
    "Figure 1 — Technology stack",
    "Figure 2 — High-level system architecture",
    "Figure 3 — Controller inheritance",
    "Figure 4 — Request lifecycle (HMVC sequence)",
    "Figure 5 — Codebase / module structure",
    "Figure 6 — Data model (Entity-Relationship)",
    "Figure 7 — Federated database connections",
    "Figure 8 — Authentication & RBAC enforcement (before hook)",
    "Figure 9 — RBAC decision model",
    "Figure 10 — Login & session flow",
    "Figure 11 — Person search to profile flow",
    "Figure 12 — Telecom data request lifecycle (core flow)",
    "Figure 13 — CDR / bulk data ingestion flow",
    "Figure 14 — Identity sync (NADRA / Verisys / Family Tree / Travel)",
    "Figure 15 — Watchlist management flow",
    "Figure 16 — Reporting & audit flow",
    "Figure 17 — Cron / background processing pipeline",
    "Figure 18 — External integrations",
    "Figure 19 — Sitemap & navigation",
    "Figure 20 — Analyst end-to-end journey",
]

doc = Document()

# ---------- base styles ----------
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)

# Heading colors
ACCENT = RGBColor(0x1F, 0x4E, 0x79)
for hname, sz in [("Heading 1", 17), ("Heading 2", 14), ("Heading 3", 12)]:
    st = doc.styles[hname]
    st.font.color.rgb = ACCENT
    st.font.size = Pt(sz)
    st.font.name = "Calibri"

def set_cell_bg(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)

def add_para(text, size=10.5, bold=False, italic=False, color=None, align=None, space_after=6):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_bullets(items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(it).font.size = Pt(10.5)

def add_figure(ordinal, width_in=6.4):
    """Insert diagram by 0-based ordinal with caption."""
    if ordinal < len(DIAGRAMS) and os.path.exists(DIAGRAMS[ordinal]):
        doc.add_picture(DIAGRAMS[ordinal], width=Inches(width_in))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(CAPTIONS[ordinal] if ordinal < len(CAPTIONS) else "")
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        cap.paragraph_format.space_after = Pt(10)
    else:
        add_para("[diagram %d not rendered]" % (ordinal + 1), italic=True,
                 color=RGBColor(0xAA, 0x00, 0x00))

def add_table(headers, rows, col_widths=None, header_bg="1F4E79"):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(hdr[i], header_bg)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci, val in enumerate(row):
            cells[ci].text = ""
            run = cells[ci].paragraphs[0].add_run(str(val))
            run.font.size = Pt(9)
            if ri % 2 == 1:
                set_cell_bg(cells[ci], "EEF3F9")
    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[ci].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t

def hrule():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), 'BBBBBB')
    pbdr.append(bottom); pPr.append(pbdr)

# =====================================================================
# COVER
# =====================================================================
for _ in range(4):
    doc.add_paragraph()
add_para("DRAMS Live", size=34, bold=True, color=ACCENT, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para("Complete System Documentation", size=18, color=RGBColor(0x44,0x44,0x44),
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
add_para("Architecture · Data Model · End-to-End Process Flowcharts",
         size=12, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x66,0x66,0x66))
for _ in range(8):
    doc.add_paragraph()
add_para("Platform:  CTD telecom-data & person-intelligence analysis system",
         size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para("Framework:  Kohana 3.x PHP (HMVC)  ·  Production base URL: https://ctd.drams.com",
         size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para("Generated:  2026-06-08", size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

# =====================================================================
# 1. EXECUTIVE OVERVIEW
# =====================================================================
doc.add_heading("1. Executive Overview", level=1)
add_para("DRAMS Live is a web-based intelligence and investigation platform used to collect, "
         "enrich, analyse and report on persons of interest and their telecommunications "
         "activity. It is built for a law-enforcement / counter-terrorism workflow and centres "
         "on five capabilities:")
add_table(["#", "Capability", "What it does"],
    [
     ["1", "Person profiling",
      "A 360 degree record per person: identity (CNIC/NADRA, foreigner, passport), addresses, education, banks, criminal record, affiliations, training, relations, devices, SIMs, photos."],
     ["2", "Telecom data acquisition",
      "Submits formal data requests (CDR, subscriber info, current location, etc.) to mobile operators via email, then automatically parses the operators' email replies back into the database."],
     ["3", "CDR analysis",
      "Call/SMS detail-record analytics: call & SMS summaries, B-party analysis, location/cell-tower tracking, device/IMEI/SIM history, graph visualisations."],
     ["4", "External data federation",
      "Searches several external government/telecom databases (subscriber, ECP, KPK-CTD, DLMS, govt-employee) plus identity services (NADRA, Verisys, Family Tree, Travel) and an AIES/CCTW REST API."],
     ["5", "Watchlists, reporting & audit",
      "Tag and watch persons, build user/person/admin analytics, and keep a full audit trail of every analyst action."],
    ],
    col_widths=[0.3, 1.7, 4.4])
add_para("The system is role-based: each user only sees the menus and data their role permits, "
         "and every search, view and request is logged.", italic=True)

# =====================================================================
# 2. TECHNOLOGY STACK
# =====================================================================
doc.add_heading("2. Technology Stack", level=1)
add_figure(0)
add_table(["Layer", "Technology"],
    [
     ["Language / Runtime", "PHP 5.x+ on Apache (XAMPP in dev)"],
     ["Framework", "Kohana 3.x (HMVC)"],
     ["ORM / DB access", "Kohana ORM, Query Builder, mysqli driver"],
     ["Auth", "Kohana auth module, ORM driver, SHA256 password hashing"],
     ["Mail", "PHPMailer (send) + PHP IMAP / webklex/php-imap (receive), Gmail OAuth"],
     ["Spreadsheets", "phpoffice/phpspreadsheet + legacy phpexcel (exports & bulk-upload parsing)"],
     ["Frontend", "Bootstrap 3.3, jQuery, DataTables, Chart.js, Select2, jQuery-Confirm, SweetAlert, Cytoscape (CDR graphs)"],
     ["Primary DB", "MySQL - aiesplus (prod) / aiesdev (dev)"],
     ["Federated DBs", "subscriber_db, ecp, ctd_kpk, DLMS_FamzSolutions (MS SQL Server), govt_emp_data"],
    ],
    col_widths=[1.8, 4.6])

# =====================================================================
# 3. HIGH-LEVEL ARCHITECTURE
# =====================================================================
doc.add_heading("3. High-Level System Architecture", level=1)
add_para("The application follows Kohana's HMVC pattern with a strict layered design. A single "
         "base controller (Controller_Working) enforces authentication and RBAC for every "
         "authenticated page; a few public controllers (login, cron, REST API) bypass it "
         "deliberately.")
add_figure(1)
doc.add_heading("Controller inheritance", level=3)
add_figure(2)

# =====================================================================
# 4. REQUEST LIFECYCLE
# =====================================================================
doc.add_heading("4. Request Lifecycle (HMVC Flow)", level=1)
add_para("Every browser hit follows the same pipeline. The default route maps "
         "/<controller>/<action>/<id> to a controller action; the default controller is login.")
add_figure(3)
doc.add_heading("Two response styles", level=3)
add_bullets([
    "Full pages render through the Kohana template wrapper (site-header, sidebar, site-footer); view bodies often live in *_functions/*.inc includes.",
    "AJAX endpoints set auto_render=FALSE and echo json_encode(...). Lists use the DataTables shape {sEcho, iTotalRecords, iTotalDisplayRecords, aaData}; charts use Chart.js arrays; the CDR network graph uses Cytoscape node/edge JSON.",
])

# =====================================================================
# 5. CODEBASE STRUCTURE
# =====================================================================
doc.add_heading("5. Codebase / Module Structure", level=1)
add_figure(4)
add_para("Include-file convention: large controllers keep each action's view-binding logic in a "
         "matching .inc include - user_functions/ (User), persons_functions/ (Persons), "
         "watchlist_functions/ (Watchlist), excel/ (report exports) and cron_job/<task>/ "
         "(per-operator send/parse logic).", italic=True)

# =====================================================================
# 6. CONTROLLER CATALOGUE
# =====================================================================
doc.add_heading("6. Controller Catalogue", level=1)
add_table(["Controller", "Auth", "Purpose"],
    [
     ["Login", "No", "Username/password + IP-throttled login, one-time token login, password recovery"],
     ["Userdashboard", "Yes", "Landing dashboard, person counts (black/grey/white), CNIC/password gating, charts"],
     ["User", "Yes", "Person/identity search, bulk search, data-upload handlers, audit utilities"],
     ["Userrequest", "Yes", "User telco data requests (CDR, subscriber, location, NADRA, family-tree, travel, blocked), scheduling, parsing queue, analytics (~60 actions)"],
     ["Userreports", "Yes", "User-activity analytics: logins, requests sent, panel/URL logs, audit & performance"],
     ["Persons", "Yes", "CDR analytics: call/SMS summaries & detail, B-party, location/cell logs, CDR graph, NADRA profile, family tree (~70 actions)"],
     ["Personprofile", "Yes", "Edit the 360 person record (info, mobiles, relations, identities, education, banks, FIRs, affiliations, trainings); Verisys/FamilyTree/Travel updates"],
     ["Personsreports", "Yes", "Person analytics: lists by category/district, devices/SIMs, top-searched, sensitive, breakups"],
     ["Adminrequest", "Yes", "Admin/DRAMS-Plus requests + bulk CSV/Excel upload, autocomplete, status management"],
     ["Adminreports", "Yes", "Identity breakups, Verisys pending/response, user breakup, blocked-IP list, menu management, password reset"],
     ["Admindatabank", "Yes", "NADRA / MSISDN databank request management & breakup reports"],
     ["Organization / Intprojects", "Yes", "Banned-organization registry / internal projects"],
     ["Watchlist", "Yes", "Tag persons, add/remove from watchlist, area-wise & user watchlists"],
     ["Databank", "Yes", "Federated search: subscriber, ECP (+family), KPK-CTD, DLMS, govt-employee"],
     ["Email / Emailtemplate / Shortcode", "Yes", "Build & send telco request emails; CRUD for templates & short-codes"],
     ["Socialanalysis / Othernumbersearch", "Yes", "Social/portal links; find other/affiliate numbers"],
     ["Upload / Download", "Yes", "File upload (CDR/IMEI/doc), MSISDN/CNIC lookups; access-controlled downloads"],
     ["Cronjob / Cronmanual", "No", "Background email send & parse, B-party rebuild, queue retries, OCR daemon, DB health"],
     ["Aiesapi", "No", "REST API for AIES/CCTW: person lookup, 3D module, fingerprints, provisioning"],
     ["Gmailapi", "No", "Gmail OAuth authorize/callback, send/read via Google API"],
     ["Verisyssync", "No", "Sync temp-uploaded Verisys/FamilyTree/Travel images into person records"],
     ["Blocked / Errors / ErrorLog / Cache", "Mixed", "Block page, error pages, error-log viewer, cache endpoints"],
    ],
    col_widths=[1.9, 0.5, 4.0])

# =====================================================================
# 7. DATA MODEL
# =====================================================================
doc.add_heading("7. Data Model (Entity-Relationship)", level=1)
add_para("The domain centres on the person. person_initiate is the identity entry point "
         "(CNIC / foreigner-ID), person is the main record, and a constellation of attribute and "
         "activity tables hang off it. Users, roles, requests and lookup tables surround the core.")
add_figure(5)
doc.add_heading("Core entity descriptions", level=3)
add_table(["Entity", "Role"],
    [
     ["person_initiate", "Identity root - CNIC / foreigner CNIC, fingerprint flag"],
     ["person", "Main record - name, father, address, region/district/police-station"],
     ["person_phone_number", "SIM-person link (owner/user, IMSI, MNC, activation/last-use)"],
     ["person_phone_device / person_device_numbers", "Devices (IMEI) and SIM-device history"],
     ["person_call_log / person_sms_log", "Raw CDR - A-party, B-party, timestamp, duration, direction, cell/location"],
     ["person_summary", "Pre-aggregated contact counts per B-party"],
     ["person_nadra_profile / person_foreigner_profile / person_detail_info", "Identity & demographic enrichment"],
     ["person attribute tables", "education, banks, criminal_record, affiliations, trainings, identities, relations, social_links"],
     ["person_tags + lu_tags", "Tagging & watchlist membership (per district, per user)"],
     ["user_request / admin_request (+nadra/familytree/travel)", "Data-acquisition workflow records"],
     ["email_messages / admin_email_messages", "Outbound/inbound email audit (linked by message_id)"],
     ["email_templates (+_type)", "Per-operator request body templates"],
     ["users / users_profile / roles / roles_users / user_tokens", "Identity & access"],
     ["banned_organizations", "Org registry referenced by affiliations"],
     ["manu_management", "Per-role menu access matrix (RBAC)"],
     ["files / inner_token", "Uploaded file tracking / stored API credentials"],
    ],
    col_widths=[2.6, 3.8])

# =====================================================================
# 8. DB CONNECTIONS
# =====================================================================
doc.add_heading("8. Database Connections (Federated Data Sources)", level=1)
add_para("database.php defines six connections. The primary (default) holds all DRAMS data; the "
         "others are read-only federated sources searched from the Databank module.")
add_figure(6)
add_table(["Connection", "Database", "Use"],
    [
     ["default", "aiesplus / aiesdev (MySQL)", "All DRAMS person, CDR, request, user data"],
     ["mobile", "subscriber_db (MySQL)", "Operator subscriber lookups"],
     ["ecp", "ecp (MySQL)", "Electoral / address & family search"],
     ["ctd_kpk", "ctd_kpk (MySQL)", "KPK CTD person profiles"],
     ["dlms_sqlsrv", "DLMS_FamzSolutions (MS SQL Server)", "Driving licenses"],
     ["govt_emp_data", "govt_emp_data (MySQL)", "Government employees"],
    ],
    col_widths=[1.5, 2.7, 2.2])

# =====================================================================
# 9. AUTH & RBAC
# =====================================================================
doc.add_heading("9. Authentication, Sessions & RBAC", level=1)
add_para("Driver: Kohana auth module, ORM driver (Auth_Orm). Passwords are SHA256-hashed. The "
         "session stores only the user_id; the full Model_User ORM object is rehydrated per "
         "request. Token lifetime is up to 14 days.")
doc.add_heading("Three login paths (Controller_Login)", level=3)
add_bullets([
    "Username/password - action_check() with per-IP brute-force throttling (~5 attempts -> IP block).",
    "One-time token - remote-workspace SSO; URL token validated for expiry, consumed once, then force_login().",
    "Password recovery - action_forget() email flow.",
])
doc.add_heading("Enforcement in Controller_Working::before()", level=3)
add_figure(7)
doc.add_heading("RBAC model", level=3)
add_para("Access is decided by combining a permission tier, a data scope, and a menu matrix:")
add_bullets([
    "Permission level - get_user_permission(user_id) returns a 0-5 tier (Administrator -> Field Officer -> Dev support).",
    "Data scope - get_user_permission_for_requests(user_id) returns all | region | district | own, filtering every request/report query.",
    "Menu access - chek_role_access(role_id, menu_id) reads the manu_management matrix to show/hide each sidebar item.",
])
add_figure(8)
add_table(["Tier", "Example roles", "Capability"],
    [
     ["1 - Administrator", "role 1", "Full access; user registration; menu management; ACL"],
     ["2 - Technical support", "roles 2, 4, 6 (HQ/Regional/District)", "Search, analysis, reports within scope"],
     ["3 - Executives", "roles 3, 5, 7", "Officer-level review within scope"],
     ["4 - Field officer", "role 8", "Limited / own scope"],
     ["5 - Dev support", "role 9", "Maintenance, menu/user admin"],
    ],
    col_widths=[1.7, 2.4, 2.3])
add_para("Additional fine-grained roles gate individual modules (Watchlist add = 25, Organization "
         "view/edit = 31/32, Databank = 34/35). Sensitive persons are further protected by "
         "cis_sensitive_person_acl.", italic=True)

# =====================================================================
# 10. CORE PROCESS FLOWCHARTS
# =====================================================================
doc.add_heading("10. Core Process Flowcharts", level=1)

doc.add_heading("10.1 Login & Session Flow", level=2)
add_figure(9)

doc.add_heading("10.2 Person Search to Profile Flow", level=2)
add_figure(10)

doc.add_heading("10.3 Telecom Data Request Lifecycle (the core flow)", level=2)
add_para("This is the heart of DRAMS: a request is created, queued, emailed to the operator, the "
         "operator emails a reply, and a cron parser writes the data back to the person record. "
         "It is fully asynchronous and email-driven.")
add_figure(11)
doc.add_heading("Status & processing codes (observed)", level=3)
add_table(["Field", "Value", "Meaning"],
    [
     ["status", "0", "pending"],
     ["status", "1", "processed / sent"],
     ["status", "2", "processing / complete-pending"],
     ["status", "3", "error"],
     ["processing_index", "3", "error"],
     ["processing_index", "4", "received"],
     ["processing_index", "5", "not found"],
     ["processing_index", "7", "uploaded"],
     ["processing_index", "8", "no data"],
    ],
    col_widths=[1.8, 0.9, 3.7])
add_para("Operators are addressed by company case in the parser switch (Mobilink/Warid, Ufone, "
         "Telenor, Zong, PTCL). Request types are keyed by user_request_type_id "
         "(1/6 = CDR, 3 = subscriber, 4 = location, 5 = NIC, 8 = Verisys, 10 = family tree, 11 = travel).",
         italic=True)

doc.add_heading("10.4 CDR / Bulk Data Ingestion Flow", level=2)
add_para("When CDR data is supplied as a file (not by email), analysts upload it directly.")
add_figure(12)

doc.add_heading("10.5 Identity Sync - NADRA / Verisys / Family Tree / Travel", level=2)
add_para("Identity-document requests follow a two-step pattern: the request goes out (email/API), "
         "images/data are dropped into a temp staging area, and Verisyssync matches them by CNIC "
         "into the person record.")
add_figure(13)

doc.add_heading("10.6 Watchlist Management Flow", level=2)
add_figure(14)

doc.add_heading("10.7 Reporting & Audit Flow", level=2)
add_figure(15)

doc.add_heading("10.8 Cron / Background Processing Pipeline", level=2)
add_para("The Cronjob controller is the asynchronous engine. Scheduled HTTP hits (or CLI) trigger "
         "send/parse/maintenance actions. Locks/cooldowns prevent overlapping IMAP runs.")
add_figure(16)

doc.add_heading("10.9 External Integrations (AIES/CCTW, Gmail, OCR)", level=2)
add_figure(17)

# =====================================================================
# 11. SITEMAP
# =====================================================================
doc.add_heading("11. Sitemap & Navigation", level=1)
add_para("The left sidebar (sidebar_user.php) is role-filtered; once a person is opened, it swaps "
         "to a person-context sidebar (sidebar_person.php).")
add_figure(18)

# =====================================================================
# 12. ROLES & JOURNEYS
# =====================================================================
doc.add_heading("12. User Roles & Journeys", level=1)
add_table(["User type", "Typical roles", "Primary journeys"],
    [
     ["Super Admin", "1", "User registration, menu management, ACL, all reports & DRAMS Plus"],
     ["Dev / Tech support", "9 (+2,4,6)", "Maintenance, user admin, system logs"],
     ["Analyst / Operator", "2-8", "Search persons, analyse CDR, submit requests, build reports"],
     ["Request manager", "14+", "Manage request queue, scheduling, projects"],
     ["Data-upload operator", "9-12", "Bulk-upload CDR files, monitor parsing"],
     ["Watchlist user", "25-26", "Tag & monitor persons"],
     ["Databank user", "34/35", "Federated identity/address searches"],
     ["Monitoring", "57-58", "Panel & URL audit logs"],
    ],
    col_widths=[1.7, 1.3, 3.4])
add_para("Representative end-to-end journey (analyst):")
add_figure(19)

# =====================================================================
# APPENDIX A
# =====================================================================
doc.add_heading("Appendix A - Database Table Inventory", level=1)
def grp(title, body):
    p = doc.add_paragraph()
    p.add_run(title + ": ").bold = True
    p.add_run(body)
grp("Person core & activity",
    "person_initiate, person, person_detail_info, person_nadra_profile, person_foreigner_profile, "
    "person_phone_number, person_phone_device, person_device_numbers, person_call_log, person_sms_log, "
    "person_summary, person_location_history, person_category, person_category_history.")
grp("Person attributes",
    "person_identities, person_education, person_banks, person_criminal_record, person_affiliations, "
    "person_trainings, person_relations, person_social_links, person_tags, person_linked_projects.")
grp("Requests & email",
    "user_request, admin_request, admin_nadra_request, admin_familytree_request, admin_travel_request, "
    "email_messages, admin_email_messages, email_templates, email_templates_type, files.")
grp("Identity / access",
    "users, users_profile, roles, roles_users, user_tokens, manu_management, cis_sensitive_person_acl, inner_token.")
grp("Lookups / geography",
    "banned_organizations, lu_tags, social_websites, telco_short_code, district, region, police_station.")
add_para("Table and column names are inferred from model/query-builder code; treat as authoritative "
         "for documentation but verify against the live schema before migrations.", italic=True,
         color=RGBColor(0x66,0x66,0x66))

# =====================================================================
# APPENDIX B
# =====================================================================
doc.add_heading("Appendix B - Glossary", level=1)
add_table(["Term", "Meaning"],
    [
     ["DRAMS", "The platform name (CTD person-intelligence & telecom-data analysis system)"],
     ["DRAMS Plus", "Admin-tier advanced/bulk request module"],
     ["CDR", "Call Detail Record - call/SMS metadata supplied by operators"],
     ["B-party", "The other party a target communicates with (called/calling number)"],
     ["MSISDN", "Mobile phone number"],
     ["IMSI / IMEI", "SIM identity / device identity"],
     ["CNIC", "Pakistani national ID number (13 digits)"],
     ["NADRA", "National identity database (identity profile source)"],
     ["Verisys", "Biometric / identity verification service"],
     ["ECP", "Electoral database (address/family search)"],
     ["DLMS", "Driving-license management system (MS SQL Server source)"],
     ["KPK-CTD", "Khyber Pakhtunkhwa CTD person database"],
     ["AIES / CCTW", "External counter-terrorism system integrating via REST API"],
     ["HMVC", "Hierarchical Model-View-Controller (Kohana's request pattern)"],
     ["manu_management", "Per-role menu-access matrix table (RBAC)"],
    ],
    col_widths=[1.6, 4.8])

hrule()
add_para("End of document - DRAMS Live System Documentation.", italic=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x88,0x88,0x88))

out = os.path.join(HERE, "DRAMS_Live_System_Documentation.docx")
doc.save(out)
print("Saved:", out)
